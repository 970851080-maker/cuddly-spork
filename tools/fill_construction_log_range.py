from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE_PATH = Path("D:/AI/电脑学习/施工日志文字.docx")
TEMPLATE_PATH = Path("D:/AI/电脑学习/施工日志.docx")
OUTPUT_PATH = Path("D:/AI/电脑学习/施工日志-3.2至5.27补全逻辑版.docx")

PROJECT_NAME = "雄安新区公共卫星测试平台项目"
START_DATE = date(2026, 3, 2)
END_DATE = date(2026, 5, 27)

WEEKDAY_NAMES = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]


@dataclass
class LogEntry:
    date: date
    weather: str
    temp_low: int | None
    temp_high: int | None
    raw_lines: list[str]
    source: str = "source"
    project_lines: list[str] = field(default_factory=list)
    unit_lines: list[str] = field(default_factory=list)
    supervision_lines: list[str] = field(default_factory=list)
    today_lines: list[str] = field(default_factory=list)
    tomorrow_lines: list[str] = field(default_factory=list)


def iter_dates(start: date, end: date) -> list[date]:
    current = start
    dates: list[date] = []
    while current <= end:
        dates.append(current)
        current += timedelta(days=1)
    return dates


def normalize_line(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def extract_docx_lines(path: Path) -> list[str]:
    doc = Document(path)
    lines: list[str] = []

    def add_text(text: str) -> None:
        text = normalize_line(text)
        if text:
            lines.append(text)

    for paragraph in doc.paragraphs:
        add_text(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    add_text(paragraph.text)
    return lines


DATE_RE = re.compile(
    r"(?P<year>20\d{2})年(?P<month>\d{1,2})月(?P<day>\d{1,2})日.*?"
    r"天气[:：](?P<weather>.*?)[，,]\s*气温[:：](?P<low>-?\d+)\s*[～~-]\s*(?P<high>-?\d+)℃"
)


def parse_source_entries(lines: list[str]) -> dict[date, LogEntry]:
    starts: list[tuple[int, date, str, int, int]] = []
    for index, line in enumerate(lines):
        match = DATE_RE.search(line)
        if not match:
            continue
        log_date = date(
            int(match.group("year")),
            int(match.group("month")),
            int(match.group("day")),
        )
        starts.append(
            (
                index,
                log_date,
                normalize_line(match.group("weather")),
                int(match.group("low")),
                int(match.group("high")),
            )
        )

    entries: dict[date, LogEntry] = {}
    for start_index, log_date, weather, low, high in starts:
        following = [candidate_index for candidate_index, *_ in starts if candidate_index > start_index]
        end_index = min(following) if following else len(lines)
        block = lines[start_index:end_index]
        entry = LogEntry(
            date=log_date,
            weather=weather,
            temp_low=low,
            temp_high=high,
            raw_lines=block,
        )
        split_entry_sections(entry)
        entries[log_date] = entry
    return entries


def section_name(line: str) -> str | None:
    if "工程名称" in line:
        return "project"
    if "施工单位" in line:
        return "unit"
    if "监理工作" in line:
        return "supervision"
    if "本日施工内容" in line:
        return "today"
    if "明日施工内容" in line:
        return "tomorrow"
    return None


def split_entry_sections(entry: LogEntry) -> None:
    buckets = {
        "project": [],
        "unit": [],
        "supervision": [],
        "today": [],
        "tomorrow": [],
    }
    active: str | None = None
    for line in entry.raw_lines[1:]:
        if line.rstrip("：:") == "一、施工时间":
            continue
        name = section_name(line)
        if name:
            active = name
        if active and not line.startswith("施工日志"):
            buckets[active].append(line)
    entry.project_lines = buckets["project"]
    entry.unit_lines = buckets["unit"]
    entry.supervision_lines = buckets["supervision"]
    entry.today_lines = buckets["today"]
    entry.tomorrow_lines = buckets["tomorrow"]


def replace_project_name(lines: list[str]) -> list[str]:
    if not lines:
        return [f"一、工程名称：{PROJECT_NAME}"]
    result = []
    for line in lines:
        if "工程名称" in line:
            result.append(f"一、工程名称：{PROJECT_NAME}")
        else:
            result.append(line)
    return result


def strip_section_heading(lines: list[str]) -> list[str]:
    return [line for line in lines if not section_name(line)]


def nearest_before(entries: dict[date, LogEntry], current: date) -> LogEntry:
    previous_dates = [item for item in entries if item < current]
    return entries[max(previous_dates)]


def nearest_after(entries: dict[date, LogEntry], current: date) -> LogEntry | None:
    next_dates = [item for item in entries if item > current]
    return entries[min(next_dates)] if next_dates else None


def infer_weather(current: date, previous: LogEntry, following: LogEntry | None) -> tuple[str, int, int]:
    if following:
        total_days = max(1, (following.date - previous.date).days)
        passed_days = (current - previous.date).days
        low = round(previous.temp_low + (following.temp_low - previous.temp_low) * passed_days / total_days)
        high = round(previous.temp_high + (following.temp_high - previous.temp_high) * passed_days / total_days)
        wind_match = re.search(r"(.风)\d+级", previous.weather)
        wind = wind_match.group(1) if wind_match else "东风"
        return f"{wind}3级，晴", low, high

    offset = (current - previous.date).days
    low = (previous.temp_low or 17) + min(3, max(0, offset // 2))
    high = (previous.temp_high or 25) + min(5, max(1, offset))
    winds = ["东风3级，晴", "东南风3级，晴", "南风3级，晴", "西南风3级，多云", "东风4级，多云", "东北风3级，晴"]
    return winds[(offset - 1) % len(winds)], low, high


def default_unit_lines(previous: LogEntry, following: LogEntry | None, current: date) -> list[str]:
    if following and (current - previous.date) > (following.date - current):
        source = following
    else:
        source = previous
    return source.unit_lines or ["二、施工单位人员及动态", "2、现场管理人员11人，机电、精装、二构等班组按施工计划组织施工。"]


def default_supervision_lines(previous: LogEntry, following: LogEntry | None) -> list[str]:
    source = previous if previous.supervision_lines else following
    return source.supervision_lines if source and source.supervision_lines else ["三、监理工作及存在问题", "监理单位巡视检查，现场按要求落实安全文明施工及质量整改。"]


def with_today_heading(lines: list[str]) -> list[str]:
    return ["四.本日施工内容：", *lines]


def with_tomorrow_heading(lines: list[str]) -> list[str]:
    return ["五.明日施工内容", *lines]


def short_text_from_adjacent(previous: LogEntry, following: LogEntry | None) -> tuple[str, str]:
    previous_text = "；".join(strip_section_heading(previous.tomorrow_lines or previous.today_lines))
    next_text = "；".join(strip_section_heading(following.today_lines)) if following else ""
    return previous_text[:90], next_text[:90]


def single_day_bridge_tasks(current: date, previous: LogEntry, following: LogEntry | None) -> list[str]:
    previous_text, next_text = short_text_from_adjacent(previous, following)
    date_hint = f"{current.month}月{current.day}日"
    return with_today_heading(
        [
            "1.机电班组：",
            f"{date_hint}承接前日计划，继续推进{previous_text or '机电安装、弱电调试及通风管线整改'}，同步复核次日施工面。",
            "2.精装及二构班组：",
            f"结合现场交叉作业，对门窗、地坪、石膏、墙板及零星砌筑部位进行收口、清理和成品保护。",
            "3.质量安全：",
            f"根据相邻日志安排，完成施工面巡查、材料堆放整理及问题整改闭合，衔接{next_text or '后续施工'}。",
        ]
    )


def tasks_0411_0423(current: date) -> list[str]:
    day = (current - date(2026, 4, 11)).days
    task_sets = [
        [
            "1.机电班组：",
            "8#首层混响室、振动间排风管及桥架继续安装，复核支吊架位置，调整局部标高。",
            "5#消防排烟风机管道、电缆敷设同步推进，检查风机盘管百叶安装尺寸。",
            "2.二构班组：材料退场、垃圾清理，步道砖破损部位修复并清扫作业面。",
            "3.精装班组：冷水间石膏施工，8号楼洁净墙板安装，室外真石漆修复及踢脚基层处理。",
        ],
        [
            "1.机电班组：",
            "8#混响室、振动间排风管完成连接段安装，桥架转角及跨接部位整理。",
            "5#消防排烟风机管道继续安装，电缆敷设后进行编号、绑扎和通道清理。",
            "2.二构班组：继续步道砖修复，配合精装班组完成施工垃圾转运。",
            "3.精装班组：冷水间石膏找平，洁净墙板收边，5号楼耐磨层局部修补。",
        ],
        [
            "1.机电班组：",
            "8#首层补风机房风管材料倒运、放线定位，旋流风口安装条件复核。",
            "二层变更风管支吊架开始安装，5#首层光纤敷设通道检查。",
            "2.二构班组：完成零星修补及现场清理，配合机电移交作业面。",
            "3.精装班组：8号楼隔声门、钢制门洞口复核，室外真石漆继续修补，踢脚线施工。",
        ],
        [
            "1.机电班组：",
            "8#补风机房风管分段安装，旋流风口连接件预装，二层变更风管同步推进。",
            "5#首层光纤开始穿线安装，整理桥架内电缆并做好标识。",
            "2.精装班组：隔声门、钢制门进场复核，8号楼首层耐磨层基层清理。",
            "3.现场管理：协调机电、精装交叉作业，完成安全文明施工检查。",
        ],
        [
            "1.机电班组：",
            "8#首层补风机房风管安装成型，旋流风口连接安装，二层变更风管调整收口。",
            "5#首层光纤安装及线路整理，检查弱电通道与风管交叉位置。",
            "2.精装班组：8号楼隔声门、钢制门安装施工，首层耐磨层分区施工并做好养护。",
            "3.质量安全：对风管连接、门框固定和耐磨层平整度进行检查整改。",
        ],
    ]
    if day < 2:
        selected = task_sets[0]
    elif day < 5:
        selected = task_sets[1]
    elif day < 8:
        selected = task_sets[2]
    elif day < 11:
        selected = task_sets[3]
    else:
        selected = task_sets[4]
    return with_today_heading(selected)


def tasks_0429_0508(current: date) -> list[str]:
    day = (current - date(2026, 4, 29)).days
    task_sets = [
        [
            "1.机电班组：",
            "8#二层变更风管、水管施工面复核，完成材料倒运、支吊架放线及局部预制。",
            "首层补风机房风管、旋流风口连接部位复查，整理5#光纤安装余量。",
            "2.精装班组：8号楼门窗安装收口，首层耐磨层养护，地坪施工区域基层清理。",
            "3.现场管理：节前检查临电、消防器材和材料堆放，落实成品保护。",
        ],
        [
            "1.机电班组：",
            "8#二层变更风管开始安装，水管支架定位，协调避让桥架及吊顶标高。",
            "二层机房、喇叭间吸音板施工条件复核，完成基层清理和排版放线。",
            "2.精装班组：耐磨层局部修补，隔声门及钢制门五金调试。",
            "3.现场管理：组织班组交叉作业协调，清理楼层通道和洞口防护。",
        ],
        [
            "1.机电班组：",
            "8#二层变更风管、水管分区安装，完成部分法兰连接及管线固定。",
            "2.吸音板班组：二层机房、喇叭间吸音板龙骨安装，复核板材排版尺寸。",
            "3.地坪施工：地坪KM6、暗室耐磨层基层处理，清理浮灰并进行边角修补。",
        ],
        [
            "1.机电班组：",
            "8#二层变更风管、水管继续安装，调整与弱电、消防管线交叉碰撞部位。",
            "2.吸音板班组：二层机房、喇叭间吸音板安装推进，完成边角收口处理。",
            "3.地坪施工：地坪KM6及暗室耐磨层施工准备，检查基层平整度和成品保护。",
        ],
        [
            "1.机电班组：",
            "8#二层变更风管、水管安装收口，检查支吊架牢固度和管线坡度。",
            "2.吸音板班组：二层机房、喇叭间吸音板继续施工，清理板面污染并修补缝隙。",
            "3.地坪施工：地坪KM6、暗室耐磨层施工，完成局部找平及养护安排。",
        ],
    ]
    if day < 2:
        selected = task_sets[0]
    elif day < 4:
        selected = task_sets[1]
    elif day < 6:
        selected = task_sets[2]
    elif day < 8:
        selected = task_sets[3]
    else:
        selected = task_sets[4]
    return with_today_heading(selected)


def tasks_0510_0512(current: date) -> list[str]:
    task_sets = {
        date(2026, 5, 10): [
            "1.机电班组：",
            "8#首层、二层变更风管安装，二层水管支架复核，整理风管连接法兰。",
            "2.吸音板班组：二层机房、喇叭间吸音板继续施工，完成局部收边。",
            "3.地坪施工：地坪KM6、暗室耐磨层施工前基层清理和标高复核。",
        ],
        date(2026, 5, 11): [
            "1.机电班组：",
            "8#二层变更风管、水管继续安装，处理与桥架、喷淋管线交叉位置。",
            "2.吸音板班组：二层机房、喇叭间吸音板安装推进，检查固定点牢固度。",
            "3.地坪施工：地坪KM6、暗室耐磨层分区施工，做好边角修补和成品保护。",
        ],
        date(2026, 5, 12): [
            "1.机电班组：",
            "8#二层变更风管、水管安装收口，进行支吊架调整和管线清理。",
            "2.吸音板班组：二层机房、喇叭间吸音板完成主要安装，局部缺陷修补。",
            "3.精装班组：电梯前室地砖维修准备，清理施工面并复核材料。",
        ],
    }
    return with_today_heading(task_sets[current])


def tasks_0514_0517(current: date) -> list[str]:
    task_sets = {
        date(2026, 5, 14): [
            "1.机电班组：",
            "8#二层变更风管、水管继续安装，弱电线缆整理，完成局部试通检查。",
            "2.精装班组：电梯前室地砖维修，二层机房墙面腻子基层处理。",
            "3.吸音板班组：二层机房、喇叭间吸音板修边，清理板面污染。",
        ],
        date(2026, 5, 15): [
            "1.机电班组：",
            "8#二层风管、水管调整收口，喇叭间强弱电穿线配管准备。",
            "2.精装班组：电梯前室地砖维修收口，气化水池上方石膏基层放线。",
            "3.二构班组：气化水池砌筑材料倒运，作业面清理并复核尺寸。",
        ],
        date(2026, 5, 16): [
            "1.机电班组：",
            "喇叭间强弱电穿线配管开始施工，弱电点位核对，通风维修改管准备。",
            "2.精装班组：气化水池上方石膏施工，二层机房腻子维修处理。",
            "3.二构班组：气化水池零星砌筑及打磨，清理周边材料。",
        ],
        date(2026, 5, 17): [
            "1.机电班组：",
            "强弱电喇叭间穿线配管继续施工，弱电调试前线路整理，通风局部改管。",
            "2.精装班组：气化水池上方石膏施工，二层机房腻子维修打磨。",
            "3.二构班组：气化水池砌筑、打磨同步推进，准备后续检查。",
        ],
    }
    return with_today_heading(task_sets[current])


def tasks_0522_0527(current: date) -> list[str]:
    task_sets = {
        date(2026, 5, 22): [
            "1.机电班组：",
            "强弱电：喇叭间穿线配管继续施工，弱电调试，整理线缆标识。",
            "通风：维修改管，检查风管连接和支吊架固定。",
            "2.精装班组：气化水池上方石膏施工，二层机房腻子维修处理。",
            "3.二构班组：气化水池砌筑施工，清理作业面。",
        ],
        date(2026, 5, 23): [
            "1.机电班组：",
            "强弱电：喇叭间穿线配管收口，弱电点位联调，处理局部故障点。",
            "通风：维修改管继续推进，完成局部风管加固。",
            "2.精装班组：气化水池上方石膏修补，二层机房腻子打磨。",
            "3.二构班组：气化水池砌筑及打磨施工，做好成品保护。",
        ],
        date(2026, 5, 24): [
            "1.机电班组：",
            "强弱电：弱电调试问题整改，喇叭间线管固定和标签完善。",
            "通风：维修改管收口，检查风口、管线碰撞及漏风隐患。",
            "2.精装班组：气化水池上方石膏施工收边，二层机房腻子维修。",
            "3.二构班组：气化水池砌筑面修整，清理砂浆和材料堆放。",
        ],
        date(2026, 5, 25): [
            "1.机电班组：",
            "强弱电：弱电系统复测，配合消防、通风专业进行联动检查。",
            "通风：维修改管复查，完成支吊架调整和风管清洁。",
            "2.精装班组：二层机房腻子打磨修补，气化水池上方石膏局部找平。",
            "3.二构班组：气化水池砌筑收口，处理边角缺陷。",
        ],
        date(2026, 5, 26): [
            "1.机电班组：",
            "强弱电：三方检查问题整改，弱电调试复核，整理竣工资料需要的点位信息。",
            "通风：维修改管后复查，处理漏风点和标高偏差。",
            "2.精装班组：气化水池上方石膏、二层机房腻子修补收尾。",
            "3.二构班组：气化水池打磨清理，配合检查整改。",
        ],
        date(2026, 5, 27): [
            "1.机电班组：",
            "强弱电：喇叭间弱电调试收口，完成线路标识、问题销项和成品保护。",
            "通风：维修改管完成复核，整理风口、支吊架及管线检查记录。",
            "2.精装班组：气化水池上方石膏、二层机房腻子完成修补清理。",
            "3.二构班组：气化水池砌筑打磨收尾，现场材料清运。",
        ],
    }
    return with_today_heading(task_sets[current])


def inferred_today_lines(current: date, previous: LogEntry, following: LogEntry | None) -> list[str]:
    if date(2026, 4, 11) <= current <= date(2026, 4, 23):
        return tasks_0411_0423(current)
    if date(2026, 4, 29) <= current <= date(2026, 5, 8):
        return tasks_0429_0508(current)
    if date(2026, 5, 10) <= current <= date(2026, 5, 12):
        return tasks_0510_0512(current)
    if date(2026, 5, 14) <= current <= date(2026, 5, 17):
        return tasks_0514_0517(current)
    if date(2026, 5, 22) <= current <= date(2026, 5, 27):
        return tasks_0522_0527(current)
    return single_day_bridge_tasks(current, previous, following)


def build_entries(source_entries: dict[date, LogEntry]) -> tuple[list[LogEntry], list[date]]:
    selected: dict[date, LogEntry] = {}
    inferred_dates: list[date] = []

    for current in iter_dates(START_DATE, END_DATE):
        source = source_entries.get(current)
        if source:
            source.project_lines = replace_project_name(source.project_lines)
            selected[current] = source
            continue

        previous = nearest_before(source_entries, current)
        following = nearest_after(source_entries, current)
        weather, low, high = infer_weather(current, previous, following)
        inferred = LogEntry(
            date=current,
            weather=weather,
            temp_low=low,
            temp_high=high,
            raw_lines=[],
            source="inferred",
            project_lines=[f"一、工程名称：{PROJECT_NAME}"],
            unit_lines=default_unit_lines(previous, following, current),
            supervision_lines=default_supervision_lines(previous, following),
            today_lines=inferred_today_lines(current, previous, following),
        )
        selected[current] = inferred
        inferred_dates.append(current)

    ordered = [selected[current] for current in iter_dates(START_DATE, END_DATE)]
    for index, entry in enumerate(ordered):
        if entry.source == "source":
            continue
        if index + 1 < len(ordered):
            next_today = strip_section_heading(ordered[index + 1].today_lines)
            entry.tomorrow_lines = with_tomorrow_heading(next_today)
        else:
            entry.tomorrow_lines = with_tomorrow_heading(
                [
                    "1.机电班组：整理调试记录，继续跟踪强弱电、通风系统问题销项。",
                    "2.精装及二构班组：完成修补部位清理、成品保护和移交前检查。",
                ]
            )
    return ordered, inferred_dates


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def remove_table(table) -> None:
    table._element.getparent().remove(table._element)


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_vertical_center(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), "center")


def format_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            set_cell_vertical_center(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10)


def set_row_widths(row, widths: tuple[float, float, float, float]) -> None:
    for cell, width in zip(row.cells, widths):
        cell.width = Cm(width)


def append_entry(table, entry: LogEntry) -> None:
    title_row = table.add_row()
    set_row_widths(title_row, (3.2, 4.0, 8.5, 2.8))
    title_cell = title_row.cells[0].merge(title_row.cells[-1])
    set_cell_text(title_cell, "施工日志", bold=True, center=True, size=14)

    date_row = table.add_row()
    set_row_widths(date_row, (3.2, 4.0, 8.5, 2.8))
    date_text = f"{entry.date.year}年{entry.date.month}月{entry.date.day}日"
    weekday = WEEKDAY_NAMES[entry.date.weekday()]
    weather_text = f"{entry.weather}，气温：{entry.temp_low}～{entry.temp_high}℃"
    set_cell_text(date_row.cells[0], date_text, center=True, size=10)
    set_cell_text(date_row.cells[1], weekday, center=True, size=10)
    set_cell_text(date_row.cells[2], weather_text, center=True, size=10)
    set_cell_text(date_row.cells[3], "编号：", center=True, size=10)

    rows = [
        ("一、工程名称：", PROJECT_NAME, True, False),
        ("二、施工单位人员及动态", "\n".join(strip_section_heading(entry.unit_lines)), True, False),
        ("三、监理工作及存在问题", "\n".join(strip_section_heading(entry.supervision_lines)), True, False),
        ("四.本日施工内容：", "\n".join(strip_section_heading(entry.today_lines)), True, False),
        ("五.明日施工内容", "\n".join(strip_section_heading(entry.tomorrow_lines)), True, False),
    ]
    for left, right, left_bold, right_bold in rows:
        row = table.add_row()
        set_row_widths(row, (3.2, 4.0, 8.5, 2.8))
        right_cell = row.cells[1].merge(row.cells[-1])
        set_cell_text(row.cells[0], left, bold=left_bold, center=True, size=10)
        set_cell_text(right_cell, right, bold=right_bold, center=False, size=10)


def save_document(entries: list[LogEntry]) -> None:
    doc = Document(TEMPLATE_PATH)
    for paragraph in list(doc.paragraphs):
        remove_paragraph(paragraph)
    for table in list(doc.tables):
        remove_table(table)

    table = doc.add_table(rows=0, cols=4)
    for entry in entries:
        append_entry(table, entry)

    format_table(table)
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    doc.save(OUTPUT_PATH)


def main() -> None:
    source_entries = parse_source_entries(extract_docx_lines(SOURCE_PATH))
    entries, inferred_dates = build_entries(source_entries)
    save_document(entries)
    print(f"输出文件：{OUTPUT_PATH}")
    print(f"日期范围：{START_DATE} 至 {END_DATE}")
    print(f"总天数：{len(entries)}")
    print(f"补齐日期数：{len(inferred_dates)}")
    print("补齐日期：" + "、".join(f"{item.month}月{item.day}日" for item in inferred_dates))


if __name__ == "__main__":
    main()
